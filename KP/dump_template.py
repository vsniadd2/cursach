# -*- coding: utf-8 -*-
"""Снять структуру шаблона ПЗ ddosControl.docx в JSON."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def find_template(kp_dir: Path) -> Path:
    for p in kp_dir.glob("*.docx"):
        if "ddos" in p.name.lower():
            return p
    raise FileNotFoundError(f"Шаблон не найден в {kp_dir}")


def dump() -> Path:
    kp_dir = Path(__file__).resolve().parent
    doc = Document(str(find_template(kp_dir)))
    out: list[dict] = []

    for i, p in enumerate(doc.paragraphs):
        out.append(
            {
                "i": i,
                "style": p.style.name if p.style else "",
                "text": p.text,
            }
        )

    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        out.append({"table": ti, "rows": rows, "row_count": len(rows), "col_count": len(rows[0]) if rows else 0})

    path = kp_dir / "_template_structure.json"
    path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
    print(f"written: {path}")
    return path


if __name__ == "__main__":
    dump()

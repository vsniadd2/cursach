# -*- coding: utf-8 -*-
"""
Генерация пояснительной записки «Экспого» на основе шаблона ПЗ ddosControl.docx
с сохранением стилей, отступов и шрифтов (форматирование run).
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from expogo_pz_content import (
    FUNC_END,
    FUNC_START,
    GLOBAL_REPLACEMENTS,
    LOGICAL_FUNCTIONS,
    PARAGRAPH_OVERRIDES,
    TABLES,
)

OUTPUT_NAME = "ПЗ Экспого.docx"
PLACEHOLDER_TEXT = "сюда фото"
FIGURE_CAPTION_RE = re.compile(r"^Рис\.\s", re.IGNORECASE)
FIGURE_CAPTION_ANYWHERE_RE = re.compile(r"Рис\.\s*[\d.]+\.\s*.+")


def find_template(kp_dir: Path) -> Path:
    for p in kp_dir.glob("*.docx"):
        if "ddos" in p.name.lower() and "экспого" not in p.name.lower():
            return p
    raise FileNotFoundError(f"Шаблон ПЗ ddosControl.docx не найден в {kp_dir}")


def set_para_preserve_format(paragraph, text: str) -> None:
    """Заменить текст, сохранив w:rPr первого run."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    set_para_preserve_format(cell.paragraphs[0], text)
    for para in cell.paragraphs[1:]:
        set_para_preserve_format(para, "")


def fill_table_rows(table, rows_data: list[list[str]]) -> None:
    """Заполнить таблицу по физическим w:tc (корректно при слиянии ячеек Word)."""
    from docx.table import _Cell

    for ri, tr in enumerate(table._tbl.findall(qn("w:tr"))):
        if ri >= len(rows_data):
            break
        row_values = rows_data[ri]
        for ci, tc in enumerate(tr.findall(qn("w:tc"))):
            if ci >= len(row_values):
                break
            set_cell_text(_Cell(tc, table), row_values[ci])


def apply_global_replacements(text: str) -> str:
    for old, new in GLOBAL_REPLACEMENTS:
        text = text.replace(old, new)
    return text


def paragraph_has_drawing(paragraph) -> bool:
    return any(True for _ in paragraph._element.iter(qn("w:drawing")))


def collect_drawing_paragraph_indices(doc: Document) -> list[int]:
    return [i for i, p in enumerate(doc.paragraphs) if paragraph_has_drawing(p)]


def set_paragraph_border(paragraph) -> None:
    """Рамка вокруг абзаца-заглушки под скриншот."""
    p_pr = paragraph._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "12")
        edge_el.set(qn("w:space"), "4")
        edge_el.set(qn("w:color"), "auto")
        p_bdr.append(edge_el)
    p_pr.append(p_bdr)


def extract_figure_caption(text: str) -> str | None:
    match = FIGURE_CAPTION_ANYWHERE_RE.search(text.replace("\n", " ").strip())
    if match:
        return match.group(0).strip()
    return None


def resolve_placeholder_target(doc: Document, drawing_index: int) -> tuple[int, str | None]:
    """
    Возвращает (индекс абзаца для рамки, подпись рисунка если была в том же абзаце).
    """
    para = doc.paragraphs[drawing_index]
    caption = extract_figure_caption(para.text)
    if caption:
        if drawing_index > 0 and not doc.paragraphs[drawing_index - 1].text.strip():
            return drawing_index - 1, caption
        return drawing_index, caption
    return drawing_index, None


def apply_photo_placeholder(paragraph) -> None:
    set_para_preserve_format(paragraph, PLACEHOLDER_TEXT)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing = Pt(200)
    set_paragraph_border(paragraph)


def apply_photo_placeholders(doc: Document, drawing_indices: list[int]) -> int:
    applied = 0
    for drawing_index in drawing_indices:
        target_index, caption = resolve_placeholder_target(doc, drawing_index)
        apply_photo_placeholder(doc.paragraphs[target_index])
        applied += 1
        if caption is not None:
            set_para_preserve_format(doc.paragraphs[drawing_index], caption)
    return applied


def remove_all_images(doc: Document) -> int:
    """Удалить w:drawing и w:pict из тела документа (подписи рисунков не трогаем)."""
    removed = 0
    body = doc.element.body
    for tag in (qn("w:drawing"), qn("w:pict")):
        for el in list(body.iter(tag)):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                removed += 1
    return removed


def build(output_name: str = OUTPUT_NAME) -> Path:
    kp_dir = Path(__file__).resolve().parent
    template = find_template(kp_dir)
    out_path = kp_dir / output_name

    shutil.copy2(template, out_path)
    doc = Document(str(out_path))

    # До правки текста: иначе inline-рисунки в тех же runs могут пропасть
    drawing_indices = collect_drawing_paragraph_indices(doc)

    override_set = set(PARAGRAPH_OVERRIDES.keys())

    for i, para in enumerate(doc.paragraphs):
        if i in override_set:
            set_para_preserve_format(para, PARAGRAPH_OVERRIDES[i])
        elif FUNC_START <= i <= FUNC_END:
            idx = i - FUNC_START
            if idx < len(LOGICAL_FUNCTIONS):
                set_para_preserve_format(para, LOGICAL_FUNCTIONS[idx])
        elif para.text.strip():
            new_text = apply_global_replacements(para.text)
            if new_text != para.text:
                set_para_preserve_format(para, new_text)

    for ti, table in enumerate(doc.tables):
        if ti >= len(TABLES):
            break
        fill_table_rows(table, TABLES[ti])

    remove_all_images(doc)
    apply_photo_placeholders(doc, drawing_indices)

    doc.save(str(out_path))
    return out_path


def verify(out_path: Path) -> list[str]:
    doc = Document(str(out_path))
    errors: list[str] = []

    if len(doc.tables) != 6:
        errors.append(f"Ожидалось 6 таблиц, найдено {len(doc.tables)}")

    for i in range(FUNC_START, FUNC_END + 1):
        text = doc.paragraphs[i].text.strip()
        if not text.startswith("–"):
            errors.append(f"Параграф {i}: нет формата списка функций")

    forbidden_patterns = [
        ("DDOS", re.compile(r"DDOS", re.I)),
        ("Swift", re.compile(r"\bSwift\b")),
        ("Xcode", re.compile(r"\bXcode\b")),
        ("React Native", re.compile(r"React\s*Native", re.I)),
        ("Expo", re.compile(r"\bExpo\b")),
        ("TypeScript", re.compile(r"\bTypeScript\b")),
        ("AsyncStorage", re.compile(r"\bAsyncStorage\b")),
    ]
    for i, para in enumerate(doc.paragraphs):
        for label, pattern in forbidden_patterns:
            if pattern.search(para.text):
                errors.append(f"Параграф {i}: найдено «{label}»")

    captions = [p.text for p in doc.paragraphs if FIGURE_CAPTION_RE.match(p.text.strip())]
    if len(captions) < 10:
        errors.append(f"Мало подписей рисунков: {len(captions)}")

    removed_check = 0
    for tag in (qn("w:drawing"), qn("w:pict")):
        removed_check += len(list(doc.element.body.iter(tag)))
    if removed_check > 0:
        errors.append(f"Остались изображения в XML: {removed_check}")

    placeholders = sum(1 for p in doc.paragraphs if p.text.strip() == PLACEHOLDER_TEXT)
    if placeholders != 23:
        errors.append(f"Ожидалось 23 заглушки «{PLACEHOLDER_TEXT}», найдено {placeholders}")

    return errors


def compare_with_template(template_path: Path, output_path: Path) -> list[str]:
    """Сверка структуры с шаблоном (параграфы, стили, таблицы)."""
    import json

    tpl_doc = Document(str(template_path))
    out_doc = Document(str(output_path))
    notes: list[str] = []

    if len(tpl_doc.paragraphs) != len(out_doc.paragraphs):
        notes.append(
            f"Число параграфов: шаблон {len(tpl_doc.paragraphs)}, результат {len(out_doc.paragraphs)}"
        )

    if len(tpl_doc.tables) != len(out_doc.tables):
        notes.append(f"Число таблиц: шаблон {len(tpl_doc.tables)}, результат {len(out_doc.tables)}")

    style_mismatches = 0
    for i, (tp, op) in enumerate(zip(tpl_doc.paragraphs, out_doc.paragraphs)):
        if (tp.style.name if tp.style else "") != (op.style.name if op.style else ""):
            style_mismatches += 1
            if style_mismatches <= 3:
                notes.append(f"Стиль параграфа {i}: {tp.style.name} != {op.style.name}")
    if style_mismatches:
        notes.append(f"Всего несовпадений стилей параграфов: {style_mismatches}")

    for ti, (tt, ot) in enumerate(zip(tpl_doc.tables, out_doc.tables)):
        if len(tt.rows) != len(ot.rows) or len(tt.columns) != len(ot.columns):
            notes.append(
                f"Таблица {ti}: размер {len(tt.rows)}x{len(tt.columns)} "
                f"vs {len(ot.rows)}x{len(ot.columns)}"
            )

    structure_path = template_path.parent / "_template_structure.json"
    if structure_path.exists():
        data = json.loads(structure_path.read_text(encoding="utf-8"))
        tpl_tables = [x for x in data if "table" in x]
        for ti, spec in enumerate(tpl_tables):
            if ti < len(out_doc.tables) and spec["row_count"] != len(out_doc.tables[ti].rows):
                notes.append(
                    f"Таблица {ti}: строк в шаблоне {spec['row_count']}, "
                    f"в файле {len(out_doc.tables[ti].rows)}"
                )

    if not notes:
        notes.append("Структура совпадает с шаблоном (параграфы, стили, таблицы).")
    return notes


if __name__ == "__main__":
    kp_dir = Path(__file__).resolve().parent
    path = build()
    issues = verify(path)
    compare_notes = compare_with_template(find_template(kp_dir), path)
    print(f"Создан файл: {path}")
    if issues:
        print("Предупреждения проверки:")
        for issue in issues:
            print(f"  - {issue}")
    else:
        print("Проверка содержания пройдена успешно.")
    print("Сверка со шаблоном:")
    for note in compare_notes:
        print(f"  - {note}")

# -*- coding: utf-8 -*-
import os
from docx import Document

KP_DIR = os.path.dirname(os.path.abspath(__file__))


def find_docx():
    for name in os.listdir(KP_DIR):
        if name.endswith(".docx") and not name.endswith(".bak"):
            return os.path.join(KP_DIR, name)
    raise FileNotFoundError("docx not found")


def set_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_cell_text(cell, new_text: str) -> None:
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], new_text)


def main():
    path = find_docx()
    doc = Document(path)

    fixes = [
        ("Главная панельScreen.tsx", "DashboardScreen.tsx"),
        ("Главная панельController.cs", "DashboardController.cs"),
        ("В папке клиентская часть/", "В папке frontend/"),
        ("В папке серверная часть:", "В папке backend:"),
        ("в папке серверная часть", "в папке backend"),
        ("в папке клиентская часть", "в папке frontend"),
        ("клиентская часть – каталог", "frontend – каталог"),
        ("серверная часть – каталог", "backend – каталог"),
        ("npm start в клиентская часть", "npm start в frontend"),
        ("dotnet run в папке серверная часть", "dotnet run в папке backend"),
        ("npm start в папке клиентская часть", "npm start в папке frontend"),
        ("клиентская часть-приложением", "клиентским приложением"),
        ("сделкий", "сделок"),
        ("картографическая библиотека Documentation", "PostgreSQL. Документация"),
        ("https://leafletjs.com/reference.html", "https://www.postgresql.org/docs/"),
        ("in-memory", "оперативной памяти"),
    ]

    for p in doc.paragraphs:
        t = p.text
        for old, new in fixes:
            if old in t:
                set_paragraph_text(p, t.replace(old, new))

    # Удалить хвост старой структуры Next.js
    remove_starts = (
        "В папке apps/web",
        "README.md –",
        "docs –",
        "apps –",
        "services –",
        "packages –",
        "В папке services/api",
        "Controllers/Parking",
        "Controllers/Bookings",
        "Hubs/Parking",
        "Infrastructure/Xpark",
        "DTOs/ApiDtos",
        "В папке packages/shared",
        "marketing-shell",
        "app-sidebar",
        "hero-section",
        "landing-секции",
        "parking-entry",
        "modern-payment",
        "sidebar-component",
        "radio-group.tsx",
        "Swagger, веб-сокеты",
        "button.tsx –",
        "input.tsx –",
        "card.tsx –",
        "label.tsx –",
        "Models/DomainRecords",
        "Services/JwtTokenService",
        "Хранятся общие TypeScript",
        "Controllers/AuthController.cs – регистрация",
    )
    for p in doc.paragraphs:
        s = p.text.strip()
        if any(s.startswith(x) or x in s for x in remove_starts):
            set_paragraph_text(p, "")

    # Источники: убрать OWASP на русский аналог при желании — оставим
    for p in doc.paragraphs:
        if p.text.strip().startswith("API – Application"):
            set_paragraph_text(p, "API – программный интерфейс приложения")
        if p.text.strip().startswith("CRUD –"):
            set_paragraph_text(p, "CRUD – создание, чтение, обновление, удаление")
        if p.text.strip().startswith("DTO –"):
            set_paragraph_text(p, "DTO – объект передачи данных")
        if p.text.strip().startswith("IDE –"):
            set_paragraph_text(p, "IDE – среда разработки")
        if p.text.strip().startswith("UI –"):
            set_paragraph_text(p, "UI – пользовательский интерфейс")
        if p.text.strip().startswith("UX –"):
            set_paragraph_text(p, "UX – удобство использования")
        if p.text.strip().startswith("QR –"):
            set_paragraph_text(p, "ПС – программное средство")

    # Таблица 15: окружение
    if len(doc.tables) > 15:
        for row in doc.tables[15].rows:
            for cell in row.cells:
                if "Икс-парк" in cell.text or "macOS" in cell.text or "localhost:3000" in cell.text:
                    set_cell_text(
                        cell,
                        "Окружение: тестирование в Expo Go и в браузере; API на http://localhost:5278; "
                        "PostgreSQL в Docker; клиент — ПК Windows 10.",
                    )

    # Таблица 15: массовая замена в ячейках
    cell_fixes = [
        ("Икс-парк", "Экспого"),
        ("XPARK", "Экспого"),
        ("парковк", "клиент"),
        ("Парковк", "Клиент"),
        ("бронирован", "сделк"),
        ("Бронирован", "Сделк"),
        ("брони", "сделки"),
        ("QR", "карточк"),
        ("Leaflet", "списк"),
        ("/app/find", "раздел «Клиенты»"),
        ("/app/dashboard", "главная панель"),
        ("Find0", "Cli0"),
        ("Book0", "Deal0"),
        ("Mkt0", "Aut0"),
        ("Pay0", "Task0"),
        ("Pts0", "Dash0"),
        ("Hist0", "Deal0"),
    ]
    if len(doc.tables) > 15:
        for row in doc.tables[15].rows:
            for cell in row.cells:
                t = cell.text
                for old, new in cell_fixes:
                    if old in t:
                        set_cell_text(cell, t.replace(old, new))

    doc.save(path)
    print("pass2 ok:", path)


if __name__ == "__main__":
    main()
